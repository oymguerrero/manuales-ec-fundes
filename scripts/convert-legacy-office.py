"""Convierte las plantillas HTML con extensión .doc/.xls a OOXML real.

Uso:
    python scripts/convert-legacy-office.py

Genera .docx/.xlsx, actualiza los enlaces HTML del sitio y elimina únicamente
los archivos fuente legacy que se convirtieron correctamente.
"""

from pathlib import Path
import re

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_DIRS = [ROOT / f"estandar-{letter}" / "templates" for letter in "abcd"]


def clean_text(node):
    return " ".join(node.get_text(" ", strip=True).split())


def provisional_language(markup):
    """Evita presentar los insumos de trabajo como estándares ya publicados."""
    replacements = {
        "Qué evalúa el F21 oficial": "Referencia de trabajo del F21",
        "F21 oficial": "versión de trabajo del F21",
        "documento oficial del estándar": "documento de trabajo de la propuesta",
        "documento oficial": "documento de trabajo",
        "glosario oficial": "glosario de trabajo",
        "estándar publicado por CONOCER": "propuesta de estándar aún no publicada",
    }
    for original, replacement in replacements.items():
        markup = markup.replace(original, replacement)
    return markup


def convert_doc(source):
    soup = BeautifulSoup(provisional_language(source.read_text(encoding="utf-8")), "html.parser")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    title = soup.find("h1")
    if title:
        paragraph = document.add_heading(clean_text(title), 0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body = soup.body or soup
    handled = {id(title)} if title else set()
    for node in body.find_all(["h1", "h2", "h3", "h4", "p", "li", "table"], recursive=True):
        if id(node) in handled or node.find_parent("table"):
            continue
        text = clean_text(node)
        if not text and node.name != "table":
            continue
        if node.name.startswith("h"):
            level = min(max(int(node.name[1]) - 1, 1), 3)
            document.add_heading(text, level=level)
        elif node.name == "li":
            document.add_paragraph(text, style="List Bullet")
        elif node.name == "p":
            document.add_paragraph(text)
        elif node.name == "table":
            rows = node.find_all("tr")
            width = max((len(row.find_all(["th", "td"], recursive=False)) for row in rows), default=1)
            table = document.add_table(rows=0, cols=width)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                cells = row.find_all(["th", "td"], recursive=False)
                target = table.add_row().cells
                for index, cell in enumerate(cells):
                    target[index].text = clean_text(cell)
                    if row_index == 0 or cell.name == "th":
                        for run in target[index].paragraphs[0].runs:
                            run.bold = True

    target = source.with_suffix(".docx")
    document.save(target)
    return target


def safe_sheet_name(name, used):
    base = re.sub(r"[\\/*?:\[\]]", " ", name).strip()[:31] or "Tabla"
    candidate = base
    number = 2
    while candidate in used:
        suffix = f" {number}"
        candidate = base[: 31 - len(suffix)] + suffix
        number += 1
    used.add(candidate)
    return candidate


def convert_xls(source):
    soup = BeautifulSoup(provisional_language(source.read_text(encoding="utf-8")), "html.parser")
    workbook = Workbook()
    workbook.remove(workbook.active)
    used = set()
    tables = soup.find_all("table")
    for index, table in enumerate(tables, 1):
        heading = table.find_previous(["h1", "h2", "h3", "h4"])
        name = clean_text(heading) if heading else f"Tabla {index}"
        sheet = workbook.create_sheet(safe_sheet_name(name, used))
        for row_index, row in enumerate(table.find_all("tr"), 1):
            cells = row.find_all(["th", "td"], recursive=False)
            for col_index, cell in enumerate(cells, 1):
                target = sheet.cell(row=row_index, column=col_index, value=clean_text(cell))
                target.alignment = Alignment(vertical="top", wrap_text=True)
                if row_index == 1 or cell.name == "th":
                    target.font = Font(bold=True, color="FFFFFF")
                    target.fill = PatternFill("solid", fgColor="193B69")
        sheet.freeze_panes = "A2"
        for column in range(1, sheet.max_column + 1):
            values = [str(sheet.cell(row=row, column=column).value or "") for row in range(1, sheet.max_row + 1)]
            sheet.column_dimensions[get_column_letter(column)].width = min(max(max(map(len, values), default=10) + 2, 12), 42)
    if not tables:
        sheet = workbook.create_sheet("Plantilla")
        sheet["A1"] = clean_text(soup)
        sheet["A1"].alignment = Alignment(wrap_text=True, vertical="top")
        sheet.column_dimensions["A"].width = 80
    target = source.with_suffix(".xlsx")
    workbook.save(target)
    return target


def update_links(mapping):
    replacements = {old.name: new.name for old, new in mapping.items()}
    for html in ROOT.rglob("*.html"):
        if "extras" in html.parts:
            continue
        original = html.read_text(encoding="utf-8")
        updated = original
        for old_name, new_name in replacements.items():
            updated = re.sub(re.escape(old_name) + r"(?!x)", new_name, updated)
        # Repara ejecuciones antiguas no idempotentes (.docxx/.xlsxx, etc.).
        updated = re.sub(r"\.docx+\b", ".docx", updated, flags=re.I)
        updated = re.sub(r"\.xlsx+\b", ".xlsx", updated, flags=re.I)
        if updated != original:
            html.write_text(updated, encoding="utf-8")


def build_checklist(path, title, entries):
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading(title, 0)
    document.add_paragraph(
        "Material de autoevaluación basado en documentos de trabajo. "
        "Revisa la versión oficial cuando el estándar sea aprobado y publicado."
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headings = ["Criterio o evidencia", "Preparado", "Por reforzar", "Notas / ubicación de evidencia"]
    for index, heading in enumerate(headings):
        table.rows[0].cells[index].text = heading
        table.rows[0].cells[index].paragraphs[0].runs[0].bold = True
    for entry in entries:
        cells = table.add_row().cells
        cells[0].text = entry
        cells[1].text = "☐"
        cells[2].text = "☐"
        cells[3].text = ""
    document.save(path)


def generate_c_checklists():
    directory = ROOT / "estandar-c" / "templates"
    route = BeautifulSoup((ROOT / "estandar-c" / "ruta-preparacion.html").read_text(encoding="utf-8"), "html.parser")
    product_entries = []
    for label in route.select('.printable-checklist label'):
        value = clean_text(label)
        if value.startswith("Producto "):
            product_entries.append(value)
    performance_entries = [
        "Desempeño 1.1 · Validar la estrategia de contenido con la persona responsable de la MiPyME",
        "Desempeño 2.1 · Generar contenido de texto con IA",
        "Desempeño 2.2 · Generar contenido de imagen con IA",
        "Desempeño 2.3 · Generar contenido de audio con IA",
        "Desempeño 2.4 · Generar contenido de video con IA",
        "Desempeño 3.1 · Publicar contenido en plataformas y documentar la implementación",
        "Desempeño 3.2 · Habilitar al personal responsable para operar el proceso",
        "Desempeño 4.1 · Presentar resultados, recomendaciones y obtener validación",
    ]
    build_checklist(directory / "checklist-productos.docx", "Checklist de productos · Ruta C", product_entries)
    build_checklist(directory / "checklist-desempenos.docx", "Checklist de desempeños · Ruta C", performance_entries)


def main():
    mapping = {}
    for directory in TEMPLATE_DIRS:
        for source in sorted(directory.glob("*.doc")):
            mapping[source] = convert_doc(source)
        for source in sorted(directory.glob("*.xls")):
            mapping[source] = convert_xls(source)
    update_links(mapping)
    generate_c_checklists()
    for source, target in mapping.items():
        if target.exists() and target.stat().st_size > 0:
            source.unlink()
            print(f"[ok] {source.relative_to(ROOT)} -> {target.name}")
    print(f"Convertidas {len(mapping)} plantillas a OOXML real.")


if __name__ == "__main__":
    main()
